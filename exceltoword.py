import streamlit as st
import pandas as pd
from docx import Document
import io

# Set the page configuration

# Title of the app
st.title("Excel to Word Converter")

# File uploader allows Excel upload
uploaded_file = st.file_uploader("Drag and drop an Excel file", type=["xlsx"], accept_multiple_files=False)

# Function to extract data from Excel
def extract_data_from_excel(uploaded_file):
    # Read the uploaded Excel file into a pandas DataFrame
    df = pd.read_excel(uploaded_file)
    return df

# Function to convert DataFrame to Word
def convert_df_to_word(df):
    doc = Document()
    
    # Add a title to the Word document
    doc.add_heading('Converted Excel Data', 0)

    # Add the DataFrame as a table to the Word document
    table = doc.add_table(rows=1, cols=len(df.columns))

    # Add the header row to the table
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = column
    
    # Add data rows to the table
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    return doc

# Check if a file has been uploaded
if uploaded_file is not None:
    # Extract data from the Excel file
    df = extract_data_from_excel(uploaded_file)
    
    # Display the DataFrame in Streamlit for the user to inspect
    st.write("Here is the data extracted from the Excel file:")
    st.dataframe(df)
    
    # Convert the extracted data to a Word document
    doc = convert_df_to_word(df)
    
    # Function to save the Word document in-memory
    def save_word_doc(doc):
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    
    # Add a download button for the Word document
    st.download_button(
        label="💾 Download Word Document",
        data=save_word_doc(doc),
        file_name="converted_file.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
